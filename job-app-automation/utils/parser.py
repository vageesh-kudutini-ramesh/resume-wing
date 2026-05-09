"""
Resume parser: extracts plain text, skills, and hyperlinks from PDF/DOCX uploads.
"""
import io
import re
import json
from typing import Tuple, List


# ── Comprehensive skills vocabulary ─────────────────────────────────────────

SKILLS_VOCAB = [
    # Programming languages
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "go",
    "golang", "rust", "kotlin", "swift", "ruby", "php", "scala", "r",
    "matlab", "perl", "shell", "bash", "powershell", "vba", "dart", "elixir",
    "haskell", "lua", "objective-c",
    # Web / frontend
    "html", "css", "react", "reactjs", "angular", "angularjs", "vue", "vuejs",
    "next.js", "nuxt", "svelte", "jquery", "bootstrap", "tailwind", "sass",
    "less", "webpack", "vite", "redux", "graphql", "rest api", "restful",
    "soap", "websocket", "oauth", "jwt",
    # Backend / frameworks
    "node.js", "nodejs", "express", "django", "flask", "fastapi", "spring",
    "spring boot", "asp.net", "laravel", "rails", "sinatra", "gin", "fiber",
    "nestjs", "grpc",
    # Databases
    "sql", "mysql", "postgresql", "postgres", "sqlite", "oracle", "mssql",
    "sql server", "mongodb", "redis", "elasticsearch", "cassandra", "dynamodb",
    "firebase", "supabase", "snowflake", "bigquery", "hive",
    # Cloud & DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
    "terraform", "ansible", "jenkins", "ci/cd", "github actions", "gitlab ci",
    "linux", "unix", "nginx", "apache", "helm", "prometheus", "grafana",
    # Data & ML / AI
    "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "tensorflow", "pytorch", "keras", "scikit-learn",
    "pandas", "numpy", "scipy", "matplotlib", "seaborn", "plotly",
    "data analysis", "data science", "data engineering", "etl", "spark",
    "hadoop", "airflow", "dbt", "power bi", "tableau", "looker",
    "hugging face", "llm", "large language model", "openai", "langchain",
    # Tools & methodologies
    "git", "github", "gitlab", "bitbucket", "jira", "confluence", "trello",
    "agile", "scrum", "kanban", "tdd", "bdd", "microservices", "monolith",
    "api design", "system design", "oop", "functional programming",
    "design patterns", "solid", "clean code", "code review",
    # Security
    "cybersecurity", "penetration testing", "soc", "siem", "vulnerability",
    "encryption", "ssl", "tls", "firewalls", "iam", "zero trust",
    # Soft skills (commonly on resumes)
    "leadership", "communication", "teamwork", "problem solving",
    "project management", "time management", "mentoring", "collaboration",
    # Domains
    "product management", "ux", "ui", "figma", "sketch", "adobe xd",
    "salesforce", "sap", "erp", "crm", "tableau",
]


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, List[str]]:
    """
    Extract text + hyperlink URIs from a PDF.

    Returns (text, hyperlinks). The hyperlinks list is kept SEPARATE from the
    text on purpose — appending them as plaintext was causing them to land
    inside the user's certifications section because that was the most-recently
    opened section when the appended block was parsed.

    Hyperlinks come from `page.get_links()` (PDF link annotations); they're
    the URLs that resumes hide behind icons or short labels like "LinkedIn".
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        link_uris: List[str] = []
        for page in doc:
            pages_text.append(page.get_text())
            try:
                for link in page.get_links() or []:
                    uri = link.get("uri")
                    if uri and uri not in link_uris:
                        link_uris.append(uri)
            except Exception:
                continue

        return "\n".join(pages_text), link_uris
    except ImportError:
        pass

    # Fallback: pdfplumber (no hyperlink support)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return ("\n".join(page.extract_text() or "" for page in pdf.pages), [])
    except Exception:
        return "", []


def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Extract text from DOCX. Hyperlinks list is currently always empty for DOCX."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs), []
    except Exception:
        return "", []


def extract_text(filename: str, file_bytes: bytes) -> Tuple[str, List[str]]:
    """Route to correct parser. Returns (text, hyperlinks)."""
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    return "", []


def extract_skills(text: str) -> List[str]:
    """
    Extract skills from resume text using vocabulary matching.
    Returns a deduplicated, sorted list of found skills.
    """
    text_lower = text.lower()
    found = set()
    for skill in SKILLS_VOCAB:
        # Match whole word / phrase
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            found.add(skill)
    return sorted(found)


def extract_contact_info(text: str) -> dict:
    """Extract email and phone from resume text."""
    email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
    phone_pattern = r"(\+?\d[\d\s\-().]{7,}\d)"

    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)

    return {
        "email": emails[0] if emails else "",
        "phone": phones[0].strip() if phones else "",
    }


def extract_name(text: str) -> str:
    """Heuristic: first non-empty line is usually the candidate name."""
    for line in text.splitlines():
        line = line.strip()
        if line and len(line.split()) <= 5 and not "@" in line:
            return line
    return ""


def parse_resume(filename: str, file_bytes: bytes) -> Tuple[str, List[str], List[str]]:
    """
    Full pipeline: extract text + skills + hyperlinks from uploaded resume.
    Returns (text, skills_list, hyperlinks_list).
    """
    text, hyperlinks = extract_text(filename, file_bytes)
    skills = extract_skills(text)
    return text, skills, hyperlinks
